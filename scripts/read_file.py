"""
Kayley's universal file reader.
Usage: python scripts/read_file.py <filepath>
Supports: PDF, DOCX, XLSX/XLS, CSV, TXT, MD, JSON, images (via Gemini vision), and more.
Outputs: extracted text to stdout, capped at 12000 chars.
"""

import sys
import os
import json

# Force UTF-8 output on Windows
if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

MAX_CHARS = 12000

def read_pdf(path: str) -> str:
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
        return "\n\n".join(text_parts)
    except Exception as e:
        # fallback to pdftotext CLI
        import subprocess
        result = subprocess.run(["pdftotext", path, "-"], capture_output=True, text=True)
        if result.returncode == 0:
            return result.stdout
        raise RuntimeError(f"pdfplumber failed ({e}), pdftotext also failed: {result.stderr}")

def read_docx(path: str) -> str:
    import docx
    doc = docx.Document(path)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    # also grab text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)

def read_xlsx(path: str) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(path, data_only=True)
    parts = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        parts.append(f"=== Sheet: {sheet_name} ===")
        for row in ws.iter_rows(values_only=True):
            row_text = " | ".join(str(v) for v in row if v is not None)
            if row_text.strip():
                parts.append(row_text)
    return "\n".join(parts)

def read_csv(path: str) -> str:
    import csv
    rows = []
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        reader = csv.reader(f)
        for row in reader:
            rows.append(" | ".join(row))
    return "\n".join(rows)

def read_image(path: str) -> str:
    """Use Gemini vision to describe/read text from an image."""
    import base64
    import urllib.request
    import urllib.error

    # Load env to get Gemini API key
    env_path = os.path.join(os.path.dirname(__file__), "..", ".env.local")
    api_key = None
    if os.path.exists(env_path):
        with open(env_path) as f:
            for line in f:
                if line.startswith("GEMINI_API_KEY="):
                    api_key = line.split("=", 1)[1].strip()
                    break
    if not api_key:
        return "[Image file detected but GEMINI_API_KEY not found — cannot use vision to read it]"

    ext = os.path.splitext(path)[1].lower()
    mime_map = {".jpg": "image/jpeg", ".jpeg": "image/jpeg", ".png": "image/png",
                ".gif": "image/gif", ".webp": "image/webp", ".bmp": "image/bmp"}
    mime = mime_map.get(ext, "image/jpeg")

    with open(path, "rb") as f:
        img_b64 = base64.b64encode(f.read()).decode()

    payload = json.dumps({
        "contents": [{
            "parts": [
                {"text": "Please read and extract ALL text visible in this image. If it's a chart or diagram, describe it. Be thorough and accurate."},
                {"inline_data": {"mime_type": mime, "data": img_b64}}
            ]
        }]
    }).encode()

    url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent?key={api_key}"
    req = urllib.request.Request(url, data=payload, headers={"Content-Type": "application/json"})
    try:
        with urllib.request.urlopen(req, timeout=30) as resp:
            data = json.loads(resp.read())
            return data["candidates"][0]["content"]["parts"][0]["text"]
    except Exception as e:
        return f"[Gemini vision call failed: {e}]"

def detect_and_read(path: str) -> str:
    if not os.path.exists(path):
        print(f"ERROR: File not found: {path}", file=sys.stderr)
        sys.exit(1)

    ext = os.path.splitext(path)[1].lower()

    # Text-native formats — just read directly
    text_exts = {".txt", ".md", ".markdown", ".log", ".csv", ".json", ".jsonl",
                 ".yaml", ".yml", ".toml", ".xml", ".html", ".htm", ".ts",
                 ".js", ".py", ".sh", ".sql", ".env", ".ini", ".cfg"}
    if ext in text_exts:
        if ext == ".csv":
            return read_csv(path)
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            return f.read()

    if ext == ".pdf":
        return read_pdf(path)

    if ext in {".docx"}:
        return read_docx(path)

    if ext in {".xlsx", ".xlsm"}:
        return read_xlsx(path)

    if ext in {".jpg", ".jpeg", ".png", ".gif", ".webp", ".bmp"}:
        return read_image(path)

    # Unknown — try file command to get mime type, then attempt text read
    import subprocess
    result = subprocess.run(["file", "--mime-type", "-b", path], capture_output=True, text=True)
    mime = result.stdout.strip()
    if mime.startswith("text/"):
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            return f.read()
    if "pdf" in mime:
        return read_pdf(path)
    if "image" in mime:
        return read_image(path)

    # Last resort: try reading as text anyway
    try:
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            content = f.read()
        if content.isprintable() or len([c for c in content if c.isprintable()]) / max(len(content), 1) > 0.8:
            return content
    except Exception:
        pass

    return f"[Cannot extract text from file type '{ext}' (mime: {mime}). File exists and is {os.path.getsize(path)} bytes.]"

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python scripts/read_file.py <filepath>", file=sys.stderr)
        sys.exit(1)
    path = sys.argv[1]
    text = detect_and_read(path)
    if len(text) > MAX_CHARS:
        text = text[:MAX_CHARS] + f"\n\n[Content truncated at {MAX_CHARS} chars — ask for more if needed]"
    print(text)
